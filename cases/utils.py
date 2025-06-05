from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
import os
from django.conf import settings
from datetime import datetime
from .models import Case

def generate_legal_document(case_id):
    """Generate a .docx file for the legal case."""
    doc = Document()
    case = Case.objects.get(id=case_id)
        # Document details
    document_date = datetime.now().strftime("%d/%m/%Y, %H:%M")
    tribunal_name = f"Tribunale Ordinario di [{getattr(case, 'court_no', '')}]"

    # Adding document date with a large font
    p = doc.add_paragraph()
    run = p.add_run(document_date)
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Adding title
    p = doc.add_paragraph()
    run = p.add_run("Formulario de DigitalMarketing")
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adding Tribunal Name as a heading
    doc.add_heading(tribunal_name, level=1)

    # Adding section titles (with bold and adjusted font size)
    titles = [
        "SEZIONE CIVILE IMMIGRAZIONE",
        "Ricorso ex art. 281 decies c.p.c.",
        "Con istanza di trattazione ex art. 127 ter c.p.c."
    ]

    for title in titles:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Adding list of Appellants
    if getattr(case, 'appellants', None) and case.appellants.exists():
        for index, appellant in enumerate(case.appellants.all(), start=1):
            appellant_text = (
                f"{index}) {getattr(appellant, 'name', '')}, nato/a a {getattr(appellant, 'birth_place', '')} "
                f"il {getattr(appellant, 'dob', '')}, CF {getattr(appellant, 'fical_code', '')}, "
                f"stato civile {getattr(appellant, 'marital_status', '')}, residente in {getattr(appellant, 'address', '')}"
            )
            doc.add_paragraph(appellant_text)

    # Adding bold "Tutti rappresentati e difesi"
    p = doc.add_paragraph()
    p.add_run("Tutti rappresentati e difesi").bold = True
    p.add_run(
        ", per mandato trasmesso telematicamente, dall’Avv. Federico Spuntarelli "
        "(C.F.: SPN FRC70L02H501E), e tutti elettivamente domiciliati presso e nel suo Studio in Roma "
        "al Via Angelico n.45, (indirizzo di posta elettronica certificata per la ricezione delle comunicazioni "
        "di Cancelleria: avv.federicospuntarelli@legalmail.it) espongono quanto segue."
    )

    # Case text sections (using appropriate formatting)
    case_text = [
        ("IN FATTO", True),
        (f"I ricorrenti discendono direttamente ed ininterrottamente da {getattr(case, 'descendant', '')} "
        f"nato/a a {getattr(case, 'des_birth_place', '')} il {getattr(case, 'des_birth_dod', '')} . (all.01).", False),
        (f"Il Sig. / la Sig.ra {getattr(case, 'descendant', '')} non si è mai naturalizzato/a cittadino argentino "
        f"così come dimostra il certificato di non naturalizzazione rilasciato dalle autorità competenti della "
        f"Repubblica argentina (all.02).", False),
        (f"Si è coniugato/a con {getattr(case, 'descendant', '')} cittadino/a {getattr(case, 'spouse_citizenship', '')}, "
        f"il {getattr(case, 'marriage_date', '')} a {getattr(case.marriage_place, 'line', '')}, "
        f"{getattr(case.marriage_place, 'state', '')}, {getattr(case.marriage_place, 'city', '')}, "
        f"{getattr(case.marriage_place, 'state', '')}, {getattr(case.marriage_place, 'zip_code', '')}, "
        f"{getattr(case.marriage_place, 'village', '')} (all.03).", False),
        (f"Lo/La stesso/a è deceduto/a il {getattr(case, 'grand_parents_dod', '')} a {getattr(case.grand_parents_dop, 'line', '')}, "
        f"{getattr(case.grand_parents_dop, 'state', '')}, {getattr(case.grand_parents_dop, 'village', '')} (all.04).", False)
    ]

    # Dynamically Adding Generations
    generation_labels = {
        1: "Prima Generazione",
        2: "Seconda Generazione",
        3: "Terza Generazione",
        4: "Quarta Generazione"
    }
    generations_dict = {}

    generations = getattr(case, 'generations', None)
    if generations and hasattr(generations, 'all'):
        for generation in generations.all():
            gen_number = getattr(generation, 'number', '')
            gen_label = generation_labels.get(gen_number, f"{gen_number}° Generazione")
            if gen_label not in generations_dict:
                generations_dict[gen_label] = []
            generations_dict[gen_label].append(getattr(generation, 'desc', ''))

    for gen_label, descriptions in generations_dict.items():
        case_text.append((f"{gen_label}:", True))  # Bold the generation title
        for desc in descriptions:
            case_text.append((f"- {desc}", False)) 

    # Adding the rest of the text
    case_text.append((
        """
        La linea di discendenza riportata dai ricorrenti nel presente ricorso trova esatto riscontro nella documentazione versata in atti.
        In relazione alla posizione degli odierni ricorrenti è documentalmente provato che la discendenza ha inizio da un avo italiano.
        È, altresì, provato che l'avo cittadino italiano ha mantenuto la cittadinanza sino alla nascita del suo discendente.
        Sussiste, infine, prova della mancata naturalizzazione argentina dell’avo, comprovata mediante attestazione rilasciata dalla competente Autorità straniera.
        La discendenza dall'avo italiano è comprovata dagli atti di stato civile di nascita e di matrimonio, in regola con la legalizzazione e muniti di traduzione ufficiale. 
        Si espone, inoltre, che né gli odierni ricorrenti, né gli ascendenti hanno mai rinunciato alla cittadinanza italiana interrompendo la catena di trasmissione della cittadinanza.
        Gli odierni ricorrenti hanno tutti i requisiti per ottenere l’accoglimento della domanda giudiziale contenuta nel presente ricorso, come analiticamente dimostrato e documentato per tabulas.
        Al fine di poter ottenere il riconoscimento giuridico della loro cittadinanza italiana dalla nascita, gli odierni ricorrenti hanno provato diverse volte di fare la richiesta di appuntamento presso il Consolato Generale d’Italia a Rosario, Argentina tramite il sistema di Prenotami come indicato per lo stesso Consolato, senza essere riusciti a prendere l’appuntamento. È fatto notorio, ed anche documentato per tabulas, che i Consolati italiani all’estero gestiscono detta tipologia di istanze con ritardi decennali. Il Consolato Italiano a Rosario ad oggi, come risulta dal sito internet ufficiale e come documentato in atti, sta trattando le domande presentate prima del 2018/2019.
        Parte ricorrente invoca la applicazione dell’art. 2 della L.241/1990 che prevede che i procedimenti di competenza delle Amministrazioni statali devono essere conclusi entro termini determinati e certi in conformità al principio di ragionevole durata del processo; l’art. 3 DPR n. 362/1994 prevede che l’amministrazione debba provvedere sulla domanda entro il termine di 730 giorni.
        Nel caso di specie sussiste una palese violazione di detta normativa e parte ricorrente è necessitata ad ottenere il richiesto riconoscimento per via giudiziale con il presente ricorso, atteso che l’irragionevole tempo di attesa equivale a lesione del diritto.
        """, False)
    )

    # Adding Clause 2 dynamically
    case_text.append((getattr(case, 'clause_2', ''), True))

    # Separator
    # case_text.append(("<><><>", False))

    # Adding the final request section
    case_text.append(("Tutto ciò premesso, i ricorrenti come sopra generalizzati, rappresentati e difesi", False))

    # Adding "RICORRONO" section with bold title
    case_text.append(("RICORRONO", True))
    case_text.append((
        "all'Ill.mo Tribunale adito affinché, ai sensi del secondo comma dell'art. 281 decies c.p.c., "
        "fissi con decreto l'udienza di comparizione delle parti, e il termine non superiore a dieci giorni prima dell'udienza per la costituzione dei resistenti.",
        False
    ))

    # Write to Word Document
    for text, is_bold in case_text:
        p = doc.add_paragraph()
        p.add_run(text).bold = is_bold

    # Adding Clausola 3 (Bold "Ministero dell'Interno")
    p = doc.add_paragraph()
    p.add_run("• Ministero dell'Interno").bold = False
    p.add_run(
        " (C.F.: 97149560589), in persona del Ministro pro tempore, con sede in Roma alla Piazza del Viminale n.1 - "
        "per legge rappresentato e difeso dall'Avvocatura Distrettuale dello Stato di "
    )
    p.add_run(f"[{getattr(case, 'clause_3', '')}]").bold = True  # Making case.clause_3 bold

    # Adding Clausola 4 (Bold "Pubblico Ministero")
    p = doc.add_paragraph()
    p.add_run("• Pubblico Ministero").bold = True
    p.add_run(
        ", in persona del Procuratore della Repubblica presso il Tribunale di"
    )
    p.add_run(f"[{getattr(case, 'clause_4', '')}]").bold = True  # Making case.clause_4 bold

    # Adding Conclusion Title (Bold "CONCLUSIONI")
    doc.add_paragraph().add_run("CONCLUSIONI").bold = True

    # Add conclusion and signature
    doc.add_paragraph(
        f"Voglia l'Ill.mo Tribunale adito, disattesa ogni contraria istanza, ritenuta la propria competenza e la sommarietà della cognizione della causa de qua:"
        f"riconoscere e dichiarare, previo accertamento della sussistenza di tutti i requisiti di legge, il possesso, sin dalla nascita, della"
        f" cittadinanza italiana in favore dei ricorrenti in forza delle circostanze documentate ed esposte nella narrativa del presente atto, da"
        f" intendersi richiamate nelle presenti conclusioni, , per l’effetto, ordinarsi al Ministero dell’Interno e/o ad ogni altra competente Autorità"
        f" amministrativa e consolare di procedere alle conseguenti iscrizioni, trascrizioni, annotazioni e comunicazioni di legge."
        f"Col favore delle spese di giudizio."
        f"Il procuratore delle parti rivolge a codesto Tribunale"
        f"ISTANZA"
        f"Di trattazione del presente ricorso ai sensi dell’articolo. 127 ter, c.p.c."
        f"Ai sensi e per gli effetti dell’art. 163 n.3 bis cpc dichiara che la presente domanda non è soggetta a condizioni di procedibilità."
        f"Con osservanza."
        f"Si dichiara che la domanda giudiziale contenuta nel presente ricorso rientra nella fascia di valore indeterminabile e che, pertanto, si" 
        f"versano [€ {getattr(case, 'total_payment', '')}] a titolo di contributo unificato per atti giudiziari."
    )
    doc.add_paragraph(f"{getattr(case, 'payment_place', '')}, {getattr(case, 'date_of_payment', '')}")

    # Adding Digital Signature
    p = doc.add_paragraph()
    run = p.add_run("Firmato digitalmente dall’Avv. Federico Spuntarelli")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save file
    file_name = f"legal_case_{getattr(case, 'id', '')}_{getattr(case, 'court_no', '')}.docx"
    file_path = os.path.join(settings.MEDIA_ROOT, "documents", file_name)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)

    return file_path

from django.core.exceptions import ObjectDoesNotExist

def get_object_or_none(model_class, **filters):
    try:
        return model_class.objects.get(**filters)
    except (model_class.DoesNotExist, ObjectDoesNotExist):
        return None

